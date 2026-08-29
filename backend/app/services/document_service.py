import os
import uuid
import re
from fastapi import UploadFile, HTTPException, status
from app.config import settings

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".doc", ".xlsx", ".xls", ".csv"}
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB

def validate_file_extension(filename: str) -> str:
    ext = os.path.splitext(filename)[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Unsupported file type '{ext}'. Allowed extensions: PDF, DOCX, XLSX, CSV."
        )
    return ext[1:]

async def save_uploaded_file(file: UploadFile) -> tuple[str, str, int]:
    ext = validate_file_extension(file.filename)
    
    # Read content to check size and empty state
    contents = await file.read()
    file_size = len(contents)
    
    if file_size == 0:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="The uploaded file is empty (0 bytes). Please upload a valid curriculum document."
        )
        
    if file_size > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"File size ({file_size / (1024*1024):.1f}MB) exceeds statutory limit of 50MB."
        )
        
    safe_filename = re.sub(r'[^a-zA-Z0-9_.-]', '_', file.filename)
    unique_name = f"{uuid.uuid4().hex[:8]}_{safe_filename}"
    file_path = os.path.join(settings.UPLOAD_DIR, unique_name)
    
    with open(file_path, "wb") as f:
        f.write(contents)
        
    return file_path, unique_name, file_size

def extract_text_and_pages(file_path: str) -> tuple[str, list[dict], list[dict], bool]:
    """
    Extracts raw text, page-by-page mapping, tables, and scanned status.
    Returns (raw_text, pages_data, tables_data, is_scanned).
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Uploaded file not found at: {file_path}")

    ext = os.path.splitext(file_path)[1].lower()
    raw_text = ""
    pages_data = []
    tables_data = []
    is_scanned = False

    if ext == ".pdf":
        try:
            import fitz
            doc = fitz.open(file_path)
            total_extracted_chars = 0
            
            for page_num, page in enumerate(doc, start=1):
                page_text = page.get_text()
                total_extracted_chars += len(page_text.strip())
                
                raw_text += f"\n--- Page {page_num} ---\n" + page_text
                pages_data.append({
                    "page_number": page_num,
                    "text": page_text,
                    "char_count": len(page_text)
                })
                
                # Check for tables using fitz find_tables if available
                if hasattr(page, "find_tables"):
                    try:
                        tabs = page.find_tables()
                        for tab in tabs:
                            tables_data.append({
                                "source_page": page_num,
                                "data": tab.extract()
                            })
                    except Exception:
                        pass

            doc.close()
            
            # Scanned PDF detection heuristic
            if len(pages_data) > 0 and (total_extracted_chars / len(pages_data)) < 30:
                is_scanned = True
        except Exception as e:
            raise RuntimeError(f"PyMuPDF failed to process PDF: {e}")

    elif ext in [".docx", ".doc"]:
        try:
            import docx
            doc = docx.Document(file_path)
            para_text = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    para_text.append(para.text.strip())
                    
            raw_text = "\n".join(para_text)
            pages_data.append({"page_number": 1, "text": raw_text, "char_count": len(raw_text)})

            for idx, table in enumerate(doc.tables):
                t_data = []
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells]
                    t_data.append(row_cells)
                tables_data.append({"source_page": 1, "table_index": idx, "data": t_data})
                raw_text += f"\n--- Table {idx+1} ---\n" + "\n".join([" | ".join(r) for r in t_data]) + "\n"
        except Exception as e:
            raise RuntimeError(f"python-docx failed to process DOCX: {e}")

    elif ext in [".xlsx", ".xls"]:
        try:
            import openpyxl
            wb = openpyxl.load_workbook(file_path, data_only=True)
            for idx, sheetname in enumerate(wb.sheetnames, start=1):
                sheet = wb[sheetname]
                sheet_text = f"\n--- Sheet: {sheetname} ---\n"
                sheet_rows = []
                for row in sheet.iter_rows(values_only=True):
                    row_cells = [str(val).strip() for val in row if val is not None and str(val).strip() != ""]
                    if row_cells:
                        row_str = " | ".join(row_cells)
                        sheet_text += row_str + "\n"
                        sheet_rows.append(row_cells)
                        
                raw_text += sheet_text
                pages_data.append({"page_number": idx, "sheet_name": sheetname, "text": sheet_text, "char_count": len(sheet_text)})
                tables_data.append({"source_page": idx, "sheet_name": sheetname, "data": sheet_rows})
        except Exception as e:
            raise RuntimeError(f"openpyxl failed to process XLSX: {e}")

    elif ext == ".csv":
        try:
            import pandas as pd
            df = pd.read_csv(file_path)
            raw_text = df.to_string()
            pages_data.append({"page_number": 1, "text": raw_text, "char_count": len(raw_text)})
            tables_data.append({"source_page": 1, "data": df.to_dict(orient="records")})
        except Exception as e:
            raise RuntimeError(f"pandas failed to process CSV: {e}")

    return raw_text, pages_data, tables_data, is_scanned
